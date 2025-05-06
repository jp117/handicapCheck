import sys
import os
import openpyxl
from docx import Document
from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import pickle
import re

load_dotenv()

SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']

def get_google_creds():
    creds = None
    if os.path.exists('token.json'):
        with open('token.json', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.json', 'wb') as token:
            pickle.dump(creds, token)
    return creds

def normalize_name(name):
    if not name:
        return ''
    # Remove leading/trailing whitespace and collapse multiple spaces/tabs
    return re.sub(r'\s+', ' ', name).strip()

def get_post_percentage_dict(sheet_id):
    creds = get_google_creds()
    service = build('sheets', 'v4', credentials=creds)
    sheet_name = 'PostPercentage'
    result = service.spreadsheets().values().get(
        spreadsheetId=sheet_id,
        range=f'{sheet_name}!A:F'
    ).execute()
    values = result.get('values', [])
    post_percentage = {}
    for row in values[1:]:
        if len(row) >= 6:
            name = normalize_name(row[0])
            pct = row[5]
            try:
                # Handle percent strings like '75%' or '100%'
                if isinstance(pct, str) and pct.strip().endswith('%'):
                    pct_val = float(pct.strip().replace('%',''))
                else:
                    pct_val = float(pct)
            except Exception:
                continue
            post_percentage[name] = pct_val
    return post_percentage

def get_roster_dict(roster_sheet_id):
    creds = get_google_creds()
    service = build('sheets', 'v4', credentials=creds)
    sheet_name = 'Sheet1'  # Assuming Sheet1 is the roster sheet
    result = service.spreadsheets().values().get(
        spreadsheetId=roster_sheet_id,
        range=f'{sheet_name}!A:E'
    ).execute()
    values = result.get('values', [])
    roster_dict = {}
    for row in values[1:]:
        if len(row) >= 5:
            member_number = row[4].strip()
            name = normalize_name(row[0])
            roster_dict[member_number] = name
    return roster_dict

if len(sys.argv) < 2:
    print("Usage: python handicap_check_report.py <spreadsheet_path>")
    sys.exit(1)

spreadsheet_path = sys.argv[1]
wb = openpyxl.load_workbook(spreadsheet_path)
sheet = wb.active

# Read member numbers from column AC (index 28)
member_numbers = []
for row in sheet.iter_rows(min_row=2):
    member_number = row[28].value  # Column AC
    if member_number:
        member_numbers.append(str(member_number).strip())

# Get Google Sheet IDs from env or prompt
sheet_id = os.getenv('GOOGLE_SHEET_ID')
if not sheet_id:
    sheet_id = input('Enter your Google Sheet ID: ').strip()
roster_sheet_id = os.getenv('ROSTER_SHEET_ID')
if not roster_sheet_id:
    roster_sheet_id = input('Enter your Roster Google Sheet ID: ').strip()

post_percentage = get_post_percentage_dict(sheet_id)

# Fetch the entire roster once and build a lookup dict
roster_dict = get_roster_dict(roster_sheet_id)

golfer_names = []
for member_number in member_numbers:
    name = roster_dict.get(member_number)
    norm_name = normalize_name(name) if name else None
    if name:
        golfer_names.append(norm_name)

report_golfers = []
no_history_golfers = []
for name in golfer_names:
    pct = post_percentage.get(name)
    if pct is not None and pct < 100:
        report_golfers.append((name, pct))
    elif pct is None:
        no_history_golfers.append(name)

base_name = os.path.basename(spreadsheet_path)
# Remove 'Golfer' and everything after it
if 'Golfer' in base_name:
    doc_base = base_name.split('Golfer')[0].rstrip()
else:
    doc_base = os.path.splitext(base_name)[0]
doc_name = f"{doc_base}- Handicap Check.docx"
dir_name = os.path.dirname(spreadsheet_path)
doc_path = os.path.join(dir_name, doc_name)

doc = Document()
doc_title = doc_name
if doc_title.lower().endswith('.docx'):
    doc_title = doc_title[:-5]
doc.add_heading(doc_title, 0)
doc.add_paragraph('Golfers with less than 100% Post Percentage:')
for name, pct in report_golfers:
    doc.add_paragraph(f"{name}: {pct:.0f}%", style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Golfers that have not played a round:')
for name in no_history_golfers:
    doc.add_paragraph(f"{name}", style='List Bullet')

doc.save(doc_path)
print(f"Report saved to {doc_path}")
